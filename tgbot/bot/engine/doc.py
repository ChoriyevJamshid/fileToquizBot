import docx
from pprint import pprint

async def get_docx_content(file_path):
    document = docx.Document(file_path)
    result = []
    try:
        for paragraph in document.paragraphs:
            if paragraph.text != '':
                result.append(paragraph.text)
    except Exception as e:
        pass
    return result


# def get_content(file_path):
#     document = docx.Document(file_path)
#     result = {}
#     counter = 1
#     for index, page in enumerate(document.paragraphs):
#         # Check if this paragraph marks a new page
#         # if page.paragraph_format.page_break_before:
#         #     print("New Page")
#         if (index + 1) % 5 == 1:
#             print(f"{counter})", page.text)
#             counter += 1
#         else:
#             print(page.text)
#
#
# get_content("matematika.docx")
